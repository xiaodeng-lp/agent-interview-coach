from __future__ import annotations

import argparse
import json
import re
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from pypdf import PdfReader


DEFAULT_SOURCE_DIR = Path(__file__).with_name("resume_materials")
DEFAULT_CACHE_PATH = Path(__file__).with_name("interview_corpus_cache.json")

SUPPORTED_SUFFIXES = {".md", ".txt", ".docx", ".pdf"}
SKIP_NAME_PREFIXES = ("~$", ".")
MAX_CHARS_PER_FILE = 120_000
CHUNK_SIZE = 2_000
CHUNK_OVERLAP = 250


@dataclass
class CorpusChunk:
    source: str
    title: str
    suffix: str
    priority: int
    chunk_index: int
    text: str


def file_priority(path: Path) -> int:
    name = path.name.lower()
    if "ai动态模拟面试官协议" in name:
        return 100
    if "ai面试背景材料" in name:
        return 98
    if "开发岗-终极" in name:
        return 95
    if "support" in name or "客服" in name or "质检" in name or "knowledge" in name or "知识库" in name:
        return 90
    if "pptagent" in name or "ppt生成agent" in name:
        return 84
    if "q lora" in name or "qlora" in name or "finetune" in name:
        return 72
    if "llamaindex" in name or "rag" in name:
        return 70
    if path.suffix.lower() == ".pdf":
        return 55
    return 60


def iter_source_files(source_dir: Path) -> Iterable[Path]:
    for path in sorted(source_dir.rglob("*")):
        if not path.is_file():
            continue
        if path.name.startswith(SKIP_NAME_PREFIXES):
            continue
        if path.suffix.lower() not in SUPPORTED_SUFFIXES:
            continue
        yield path


def normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def read_text_file(path: Path) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(errors="ignore")


def read_docx(path: Path) -> str:
    document = Document(str(path))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def read_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    pages: list[str] = []
    for index, page in enumerate(reader.pages):
        try:
            page_text = page.extract_text() or ""
        except Exception as exc:  # PDF extraction can fail on individual pages.
            page_text = f"[第 {index + 1} 页提取失败: {exc}]"
        if page_text.strip():
            pages.append(f"\n[PDF 第 {index + 1} 页]\n{page_text}")
    return "\n".join(pages)


def read_supported_file(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".md", ".txt"}:
        return read_text_file(path)
    if suffix == ".docx":
        return read_docx(path)
    if suffix == ".pdf":
        return read_pdf(path)
    raise ValueError(f"Unsupported file type: {path}")


def chunk_text(text: str) -> list[str]:
    text = normalize_text(text)[:MAX_CHARS_PER_FILE]
    if not text:
        return []
    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = min(start + CHUNK_SIZE, len(text))
        chunks.append(text[start:end].strip())
        if end == len(text):
            break
        start = max(0, end - CHUNK_OVERLAP)
    return [chunk for chunk in chunks if chunk]


def build_corpus(source_dir: Path = DEFAULT_SOURCE_DIR) -> list[CorpusChunk]:
    corpus: list[CorpusChunk] = []
    for path in iter_source_files(source_dir):
        try:
            text = read_supported_file(path)
        except Exception as exc:
            text = f"[文件读取失败: {exc}]"
        for index, chunk in enumerate(chunk_text(text)):
            corpus.append(
                CorpusChunk(
                    source=str(path),
                    title=path.name,
                    suffix=path.suffix.lower(),
                    priority=file_priority(path),
                    chunk_index=index,
                    text=chunk,
                )
            )
    corpus.sort(key=lambda item: (-item.priority, item.title, item.chunk_index))
    return corpus


def save_corpus_cache(corpus: list[CorpusChunk], cache_path: Path = DEFAULT_CACHE_PATH) -> None:
    payload = {
        "source_dir": str(DEFAULT_SOURCE_DIR),
        "chunk_count": len(corpus),
        "files": sorted({chunk.source for chunk in corpus}),
        "chunks": [asdict(chunk) for chunk in corpus],
    }
    cache_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def load_corpus_cache(cache_path: Path = DEFAULT_CACHE_PATH) -> list[CorpusChunk]:
    payload = json.loads(cache_path.read_text(encoding="utf-8"))
    return [CorpusChunk(**item) for item in payload["chunks"]]


def select_context(corpus: list[CorpusChunk], user_message: str, max_chars: int = 18_000) -> str:
    keywords = set(re.findall(r"[\w\u4e00-\u9fff]{2,}", user_message.lower()))

    def score(chunk: CorpusChunk) -> tuple[int, int]:
        text = (chunk.title + "\n" + chunk.text[:800]).lower()
        keyword_hits = sum(1 for keyword in keywords if keyword in text)
        return chunk.priority + keyword_hits * 8, chunk.priority

    selected: list[CorpusChunk] = []
    used = 0
    for chunk in sorted(corpus, key=score, reverse=True):
        block = f"\n\n### {chunk.title} / chunk {chunk.chunk_index}\n{chunk.text}"
        if used + len(block) > max_chars:
            continue
        selected.append(chunk)
        used += len(block)
        if used >= max_chars:
            break
    return "".join(
        f"\n\n### {chunk.title} / chunk {chunk.chunk_index}\n{chunk.text}" for chunk in selected
    ).strip()


def main() -> None:
    parser = argparse.ArgumentParser(description="Build interview context corpus from resume materials.")
    parser.add_argument("--source-dir", type=Path, default=DEFAULT_SOURCE_DIR)
    parser.add_argument("--cache", type=Path, default=DEFAULT_CACHE_PATH)
    args = parser.parse_args()

    corpus = build_corpus(args.source_dir)
    save_corpus_cache(corpus, args.cache)
    files = sorted({chunk.source for chunk in corpus})
    print(f"Indexed {len(files)} files into {len(corpus)} chunks.")
    print(f"Cache: {args.cache}")
    for file in files:
        print(f"- {file}")


if __name__ == "__main__":
    main()
